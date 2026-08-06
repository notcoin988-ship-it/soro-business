"""Генератор документов для тестов парсеров (раздел 6.2 ТЗ).

Зачем синтетика, когда есть настоящие документы банка (лежат в
`tests/data`, проверяются в `test_parsers_real.py`): здесь фикстура
собирается ровно под одну ветку кода — короткая страница, скан без
текстового слоя, таблица с известным содержимым. В настоящем документе
такую страницу ещё надо найти, а скана среди PDF Эсхаты не оказалось
вовсе. Содержимое взято близко к настоящим тарифам: таджикский с
диакритикой (ӣ ӯ ҳ ҷ ғ қ), русский, проценты, суммы в сомонӣ, таблицы.

Файлы создаются во временной папке теста и в git не попадают: генератор
всегда собирает актуальную версию. Настоящие документы, наоборот, лежат
в git — банк их правит и удаляет, а тест обязан быть воспроизводимым.

Отдельный сорт фикстуры — **скан**: страница без текстового слоя, только
картинка. На ней проверяется OCR-ветка (< 40 символов → pytesseract,
lang='tgk+rus').
"""

from __future__ import annotations

from pathlib import Path

# Шрифт нужен обеим сторонам: без него PDF не покажет кириллицу, а скан
# нечего будет распознавать. В образе backend стоит dejavu (см. Dockerfile),
# на Windows — обычные системные шрифты.
FONT_CANDIDATES = (
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
    "C:/Windows/Fonts/arial.ttf",
    "C:/Windows/Fonts/segoeui.ttf",
)

# Текст первой страницы: то, ради чего всё и делается, — таджикская
# банковская страница с диакритикой и числами.
PDF_PAGE_TJ = (
    "Тарифҳои амонати «Ояндасоз»\n"
    "Фоизи солона: 14,5%\n"
    "Маблағи ҳадди ақал: 500 сомонӣ\n"
    "Мӯҳлат: аз 12 то 36 моҳ\n"
    "Ҷуброни пеш аз мӯҳлат: бо фоизи 0,5%"
)
PDF_PAGE_RU = (
    "Условия обслуживания\n"
    "Выпуск зарплатной карты — бесплатно\n"
    "Лимит снятия наличных: 10 000 сомони в сутки\n"
    "Комиссия за перевод внутри банка: 0%"
)
# Страница-разделитель: меньше 40 символов, текстовый слой есть.
# Проверяет, что порог OCR смотрит именно на длину, а не на «есть текст».
PDF_PAGE_SHORT = "Замима 1"

SCAN_TEXT_LINES = (
    "Тарифҳои бонк",
    "Фоизи солона 14,5 фоиз",
    "Маблағи ҳадди ақал 500 сомонӣ",
)


def find_font() -> str | None:
    for path in FONT_CANDIDATES:
        if Path(path).exists():
            return path
    return None


def _require_font() -> str:
    font = find_font()
    if font is None:  # pragma: no cover — на машинах без шрифтов тесты skip
        raise RuntimeError("не найден TTF-шрифт с кириллицей")
    return font


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def make_pdf(path: Path, pages: list[str], fontsize: int = 13) -> Path:
    """PDF с текстовым слоем.

    Текст кладётся через `insert_textbox`, а не `insert_text`: последний не
    переносит строки, и всё, что не влезло по ширине или высоте, молча
    обрезается. На замере это давало вчетверо меньше символов на странице,
    чем задумано, — и вчетверо более оптимистичную цифру индексации.
    """
    import fitz

    font = _require_font()
    doc = fitz.open()
    for content in pages:
        page = doc.new_page()
        box = page.rect + (56, 56, -56, -56)  # поля 2 см
        page.insert_textbox(
            box, content, fontsize=fontsize, fontfile=font, fontname="tj"
        )
    doc.save(str(path))
    doc.close()
    return path


def make_scan_pdf(path: Path, lines: tuple[str, ...] = SCAN_TEXT_LINES) -> Path:
    """PDF-скан: страница отрисована в картинку, текстового слоя нет.

    Ровно то, что приходит от банка «отсканировали приказ и прислали».
    Кегль крупный и dpi высокий — иначе tesseract на таджикской диакритике
    начинает фантазировать, а тест должен проверять ветку кода, а не
    качество распознавания.
    """
    import fitz

    font = _require_font()

    text_doc = fitz.open()
    page = text_doc.new_page()
    y = 100
    for line in lines:
        page.insert_text((60, y), line, fontsize=26, fontfile=font, fontname="tj")
        y += 50
    pixmap = page.get_pixmap(dpi=200)
    text_doc.close()

    scan = fitz.open()
    image_page = scan.new_page(width=pixmap.width * 0.75, height=pixmap.height * 0.75)
    image_page.insert_image(image_page.rect, pixmap=pixmap)
    scan.save(str(path))
    scan.close()
    return path


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

DOCX_TABLE = [
    ["Валюта", "Купля", "Продажа"],
    ["USD", "10,90", "11,05"],
    ["EUR", "11,80", "12,00"],
]


def make_docx(path: Path) -> Path:
    """DOCX: абзац, таблица, снова абзац — порядок важен для теста."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Қоидаҳои хизматрасонӣ")
    doc.add_paragraph("Курси асъор дар санаи 04.08.2026:")

    table = doc.add_table(rows=0, cols=3)
    for row in DOCX_TABLE:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value

    doc.add_paragraph("")  # пустой абзац: должен выпасть из текста
    doc.add_paragraph("Маълумоти иловагӣ бо телефони 44 600 60 60.")
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------


def make_xlsx(path: Path) -> Path:
    """XLSX: два листа, пустая строка внутри данных, целое число как float."""
    from openpyxl import Workbook

    book = Workbook()

    rates = book.active
    rates.title = "Фоизҳо"
    rates.append(["Маҳсулот", "Фоиз", "Ҳадди ақал"])
    rates.append(["Ояндасоз", "14,5%", 500])
    rates.append([None, None, None])  # пустая строка — пропускается
    rates.append(["Пасандози кӯдакона", "12%", 200])

    fees = book.create_sheet("Комиссияҳо")
    fees.append(["Амалиёт", "Комиссия"])
    fees.append(["Интиқол дар дохили бонк", "0%"])
    fees.append(["Интиқол ба бонки дигар", "0,5%"])

    book.save(str(path))
    return path


# ---------------------------------------------------------------------------
# 40-страничный PDF для замера индексации (критерий сдачи недели 2)
# ---------------------------------------------------------------------------


def make_big_pdf(path: Path, pages: int = 40, density: int = 1) -> Path:
    """PDF на `pages` страниц с текстом банковского вида.

    Норматив ТЗ раздела 6 — «40-страничный PDF индексируется < 3 минут на
    CPU». Меряется на этом файле, см. scripts/bench_index.py.

    `density` — сколько раз повторить блок на странице. При density=1
    получается ~180 символов на страницу; это НЕ похоже на настоящий тариф,
    где страница плотно набита (2000–3000 символов), а от объёма текста
    напрямую зависит число чанков и, значит, время эмбеддингов. Для
    честного замера нужен density≈14.
    """
    blocks = [PDF_PAGE_TJ, PDF_PAGE_RU]
    content = []
    for i in range(1, pages + 1):
        block = blocks[i % len(blocks)]
        body = "\n".join(
            f"Банди {i}.{k} — " + block.replace("\n", " ")
            for k in range(1, density + 1)
        )
        content.append(f"Саҳифаи {i}\n{body}")
    # кегль 8: при 13 на страницу влезает ~1000 знаков, а надо 2500 — как в
    # настоящем тарифе, набранном мелким шрифтом в две трети листа
    return make_pdf(path, content, fontsize=8)


def make_scan_book(path: Path, pages: int = 5) -> Path:
    """Многостраничный скан — для замера цены OCR-ветки.

    OCR идёт постранично и на CPU стоит секунды на страницу, поэтому
    40-страничный скан ведёт себя совсем не так, как 40-страничный
    текстовый PDF. Меряем на нескольких страницах и пересчитываем.
    """
    import fitz

    font = _require_font()

    scan = fitz.open()
    for number in range(1, pages + 1):
        text_doc = fitz.open()
        page = text_doc.new_page()
        y = 90
        lines = (f"Саҳифаи {number}", *SCAN_TEXT_LINES, *SCAN_TEXT_LINES)
        for line in lines:
            page.insert_text((55, y), line, fontsize=22, fontfile=font, fontname="tj")
            y += 44
        pixmap = page.get_pixmap(dpi=200)
        text_doc.close()

        image_page = scan.new_page(
            width=pixmap.width * 0.75, height=pixmap.height * 0.75
        )
        image_page.insert_image(image_page.rect, pixmap=pixmap)

    scan.save(str(path))
    scan.close()
    return path
