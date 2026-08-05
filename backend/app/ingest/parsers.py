"""Извлечение текста из документов базы знаний (раздел 6.2 ТЗ).

Один вход — `parse_file(path)`, один выход — список `ParsedPage`. Дальше по
пайплайну идёт чанкер (6.3), поэтому парсер НЕ режет текст и НЕ дописывает
контекст-шапку: его задача — отдать текст постранично и сказать, что считать
страницей.

Что такое «страница» для каждого типа (поле `chunks.page` в DDL раздела 5):

* PDF  — номер страницы, с единицы;
* XLSX — номер листа, с единицы (правило ТЗ: «page = номер листа»);
* DOCX — страниц нет: python-docx работает с потоком абзацев, а разбивка на
  страницы появляется только при вёрстке в Word. Отдаём одну ParsedPage с
  `page=None`;
* HTML — см. `parse_html`, страницы нет, номер проставляет краулер.

ОТКЛОНЕНИЕ ОТ ТЗ: в DDL `chunks.page INT`,
а комментарий рядом — «страница/лист/URL-хвост». URL-хвост в INT не влезает.
Здесь это не решается, парсер просто отдаёт `page=None` для веба.
"""

from __future__ import annotations

import io
import zipfile
from dataclasses import dataclass
from pathlib import Path

# Порог из ТЗ: страница короче 40 символов считается сканом и идёт в OCR.
OCR_MIN_CHARS = 40
# Языки OCR из ТЗ. Пакеты tesseract-ocr-tgk и -rus стоят в Dockerfile.
OCR_LANG = "tgk+rus"
# РЕШЕНИЕ (в ТЗ не задано): плотность рендера страницы перед OCR.
# 300 dpi — стандарт для распознавания текста; на 150 tesseract заметно
# хуже берёт таджикскую диакритику (ӣ ӯ ҳ ҷ ғ қ).
OCR_DPI = 300

SUPPORTED_SUFFIXES = {".pdf": "pdf", ".docx": "docx", ".xlsx": "xlsx"}


@dataclass(frozen=True)
class ParsedPage:
    """Одна страница/лист документа. `page` попадает в `chunks.page`."""

    page: int | None
    text: str
    # True, если текст получен распознаванием, а не из текстового слоя.
    # Нужно на приёмке: OCR-текст грязнее, и это видно в логе индексации.
    ocr: bool = False


class UnsupportedDocument(Exception):
    """Расширение файла вне списка pdf/docx/xlsx (CHECK `documents.kind`)."""


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def parse_pdf(path: str | Path) -> list[ParsedPage]:
    """PDF через pymupdf, постранично `page.get_text('text')`.

    Если на странице меньше `OCR_MIN_CHARS` символов — считаем страницу
    сканом и гоним через pytesseract с `lang='tgk+rus'`.

    Импорты внутри функции намеренно: pymupdf и pytesseract тяжёлые, а
    parsers импортируется в том числе тестами краулера, которым PDF не нужен.
    """
    import fitz  # pymupdf

    pages: list[ParsedPage] = []
    with fitz.open(path) as doc:
        for number, page in enumerate(doc, start=1):
            text = page.get_text("text").strip()
            if len(text) < OCR_MIN_CHARS:
                ocr_text = _ocr_page(page)
                # если OCR тоже ничего не дал — страница действительно пустая
                # (титульный лист с одной картинкой-логотипом, разделитель),
                # и в базу она уйдёт пустой строкой, а не потеряется
                pages.append(ParsedPage(page=number, text=ocr_text, ocr=True))
            else:
                pages.append(ParsedPage(page=number, text=text))
    return pages


def _ocr_page(page) -> str:
    """Рендер страницы в картинку и распознавание. Ошибку не глушим."""
    import pytesseract
    from PIL import Image

    pix = page.get_pixmap(dpi=OCR_DPI)
    image = Image.open(io.BytesIO(pix.tobytes("png")))
    return pytesseract.image_to_string(image, lang=OCR_LANG).strip()


def ocr_available() -> bool:
    """Есть ли в системе tesseract с нужными языками.

    Используется тестами: вне контейнера tesseract обычно не стоит, и тест
    OCR-ветки надо не падать, а пропускать.
    """
    try:
        import pytesseract

        langs = set(pytesseract.get_languages())
    except Exception:
        return False
    return {"tgk", "rus"} <= langs


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

# РЕШЕНИЕ (трактовка ТЗ): «таблицу превращать в строки вида «Ячейка1 |
# Ячейка2», заголовок таблицы повторять в каждой строке». Повторять шапку
# нужно потому, что чанкер режет текст по абзацам и строка таблицы легко
# оторвётся от своей шапки — в отрыве «14,5 | 500» бессмысленно. Формат:
#   Валюта: USD | Курс покупки: 10.90
# то есть заголовок приклеен к своей ячейке. Это совпадает с правилом для
# XLSX («Заголовок_колонки: значение; ...») и даёт одинаковый вид таблиц
# из обоих форматов в поиске.
CELL_SEP = " | "


def parse_docx(path: str | Path) -> list[ParsedPage]:
    """DOCX через python-docx: абзацы + таблицы в порядке следования.

    python-docx не отдаёт перемешанный поток «абзац/таблица», поэтому идём
    по дочерним элементам тела документа сами: иначе все таблицы уедут в
    конец текста и потеряют привязку к своему разделу.
    """
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    document = Document(str(path))
    body = document.element.body
    lines: list[str] = []

    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            text = Paragraph(child, document).text.strip()
            if text:
                lines.append(text)
        elif tag == "tbl":
            lines.extend(_docx_table_lines(Table(child, document)))

    return [ParsedPage(page=None, text="\n".join(lines))]


def _docx_table_lines(table) -> list[str]:
    """Таблица → строки «Заголовок: значение | Заголовок: значение»."""
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    rows = [row for row in rows if any(row)]
    if not rows:
        return []

    header, *data = rows
    if not data:
        # таблица из одной строки: шапки нет, отдаём как есть
        return [CELL_SEP.join(c for c in header if c)]

    lines: list[str] = []
    for row in data:
        parts = []
        for index, cell in enumerate(row):
            if not cell:
                continue
            title = header[index] if index < len(header) else ""
            parts.append(f"{title}: {cell}" if title else cell)
        if parts:
            lines.append(CELL_SEP.join(parts))
    return lines


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------


def parse_xlsx(path: str | Path) -> list[ParsedPage]:
    """XLSX через openpyxl: каждый лист отдельно, `page` = номер листа.

    Строка превращается в «Заголовок_колонки: значение; ...», пустые строки
    пропускаются (правило 6.2).
    """
    from openpyxl import load_workbook

    # data_only=True — нужны посчитанные значения формул, а не сами формулы:
    # в тарифах банка проценты часто вычисляются. Файл без кеша значений
    # отдаст None — это ожидаемо и попадает под «пустую ячейку».
    workbook = load_workbook(str(path), read_only=True, data_only=True)
    pages: list[ParsedPage] = []
    try:
        for index, sheet in enumerate(workbook.worksheets, start=1):
            lines = _xlsx_sheet_lines(sheet)
            if lines:
                pages.append(ParsedPage(page=index, text="\n".join(lines)))
    finally:
        workbook.close()
    return pages


def _xlsx_sheet_lines(sheet) -> list[str]:
    header: list[str] = []
    lines: list[str] = []

    for row in sheet.iter_rows(values_only=True):
        values = [_cell_to_text(v) for v in row]
        if not any(values):
            continue
        if not header:
            # первая непустая строка листа — шапка
            header = values
            continue
        parts = [
            f"{header[i]}: {value}" if i < len(header) and header[i] else value
            for i, value in enumerate(values)
            if value
        ]
        if parts:
            lines.append("; ".join(parts))

    return lines


def _cell_to_text(value) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        # openpyxl отдаёт целые числа как 500.0 — в тексте это мусор
        return str(int(value))
    return str(value).strip()


# ---------------------------------------------------------------------------
# HTML — используется краулером (6.2, ветка «Сайт»)
# ---------------------------------------------------------------------------

# Выкидываем по ТЗ: nav, header, footer, script. Плюс style/noscript/svg —
# РЕШЕНИЕ: это тот же служебный мусор, в тексте от них только шум.
DROP_TAGS = ("nav", "header", "footer", "script", "style", "noscript", "svg")
# Приоритет контейнеров из ТЗ: main → article → body.
CONTENT_TAGS = ("main", "article", "body")


def parse_html(html: str) -> tuple[str, str]:
    """HTML → (заголовок, текст). Возвращает пару, потому что title страницы
    идёт в `documents.title`, а текст — в чанки.

    РЕШЕНИЕ (в ТЗ не сказано, откуда брать title): <title>, если пуст —
    первый <h1>, если и его нет — пустая строка, и краулер подставит
    хвост URL.
    """
    from selectolax.parser import HTMLParser

    tree = HTMLParser(html)

    title = ""
    if tree.css_first("title") is not None:
        title = _collapse(tree.css_first("title").text())
    if not title and tree.css_first("h1") is not None:
        title = _collapse(tree.css_first("h1").text())

    for tag in DROP_TAGS:
        for node in tree.css(tag):
            node.decompose()

    node = None
    for tag in CONTENT_TAGS:
        node = tree.css_first(tag)
        if node is not None:
            break
    if node is None:
        return title, ""

    # separator='\n' — иначе соседние блоки слипаются в одно слово и
    # чанкеру нечего резать: он работает по абзацам
    text = node.text(separator="\n", strip=True)
    return title, _clean_lines(text)


def _collapse(value: str) -> str:
    return " ".join(value.split())


def _clean_lines(text: str) -> str:
    lines = [_collapse(line) for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


# ---------------------------------------------------------------------------
# диспетчер
# ---------------------------------------------------------------------------


def parse_file(path: str | Path) -> list[ParsedPage]:
    """Разбирает файл по расширению. Тип — как в `documents.kind`."""
    path = Path(path)
    kind = SUPPORTED_SUFFIXES.get(path.suffix.lower())
    if kind is None:
        raise UnsupportedDocument(
            f"{path.name}: поддерживаются только pdf/docx/xlsx"
        )
    if kind == "pdf":
        return parse_pdf(path)
    if kind == "docx":
        _assert_zip(path, "docx")
        return parse_docx(path)
    _assert_zip(path, "xlsx")
    return parse_xlsx(path)


def _assert_zip(path: Path, kind: str) -> None:
    """docx/xlsx — это zip-архивы. Файл с правильным расширением и чужим
    содержимым (переименованный doc, обрезанная загрузка) роняет библиотеку
    невнятной ошибкой; отвечаем понятной, она уйдёт в `documents.error`.
    """
    if not zipfile.is_zipfile(path):
        raise UnsupportedDocument(
            f"{path.name}: расширение .{kind}, но это не {kind}-файл"
        )
